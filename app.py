from flask import Flask, request, render_template_string, send_file
from docx import Document
import subprocess
import os
import math

app = Flask(__name__)

def compute_skus(num_leases):
    """Return SKUs based on number of leases."""
    try:
        n = int(num_leases)
    except:
        n = 0
    if n < 0:
        n = 0

    skus = []

    # Always include PS-TECH
    skus.append({
        'sku': 'PS-TECH',
        'hours': 16,
        'rate': 275,
        'total': 4400
    })

    # IFM SKU selection
    if n <= 15:
        skus.append({
            'sku': 'IFM-LEASE-BASE',
            'description': 'Per CPQ',
            'quantity': 1
        })
    elif n <= 75:
        skus.append({
            'sku': 'IFM-LEASE-STAN',
            'description': 'Per CPQ',
            'quantity': 1
        })
    elif n <= 150:
        skus.append({
            'sku': 'IFM-LEASE-PROF',
            'description': 'Per CPQ',
            'quantity': 1
        })
    else:
        skus.append({
            'sku': 'IFM-LEASE-PROF',
            'description': 'Per CPQ',
            'quantity': 1
        })
        additional = n - 150
        packs = math.ceil(additional / 50)
        skus.append({
            'sku': 'IFM-LEASE-PACK',
            'description': 'Per CPQ (per 50 additional leases)',
            'quantity': packs
        })

    return skus

# Sage brand colors
# Primary Green: #00DC00 (Sage green)
# Dark Green: #1D6F42
# Text Dark: #1E1E1E
# Background: #F7F7F7

HTML_FORM = """
<!DOCTYPE html>
<html>
<head>
    <title>SOW Generator | Sage</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            background: #F7F7F7;
            color: #1E1E1E;
            min-height: 100vh;
        }
        
        .header {
            background: #FFFFFF;
            padding: 20px 40px;
            border-bottom: 1px solid #E5E5E5;
            display: flex;
            align-items: center;
            gap: 12px;
        }
        
        .logo {
            width: 80px;
            height: auto;
        }
        
        .header-text {
            font-size: 24px;
            font-weight: 700;
            color: #1D6F42;
        }
        
        .container {
            max-width: 600px;
            margin: 60px auto;
            padding: 0 20px;
        }
        
        h1 {
            font-size: 32px;
            font-weight: 700;
            color: #1E1E1E;
            margin-bottom: 8px;
        }
        
        .subtitle {
            color: #666666;
            font-size: 16px;
            margin-bottom: 40px;
        }
        
        .form-card {
            background: #FFFFFF;
            padding: 40px;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
        }
        
        .form-group {
            margin-bottom: 24px;
        }
        
        label {
            display: block;
            font-size: 14px;
            font-weight: 600;
            color: #1E1E1E;
            margin-bottom: 8px;
        }
        
        input {
            width: 100%;
            padding: 14px 16px;
            font-size: 16px;
            border: 2px solid #E5E5E5;
            border-radius: 8px;
            transition: border-color 0.2s, box-shadow 0.2s;
            font-family: inherit;
        }
        
        input:focus {
            outline: none;
            border-color: #00DC00;
            box-shadow: 0 0 0 3px rgba(0, 220, 0, 0.15);
        }
        
        input::placeholder {
            color: #AAAAAA;
        }
        
        .btn-primary {
            width: 100%;
            background: #00DC00;
            color: #FFFFFF;
            padding: 16px 24px;
            font-size: 16px;
            font-weight: 600;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            transition: background 0.2s, transform 0.1s;
            font-family: inherit;
        }
        
        .btn-primary:hover {
            background: #00C400;
        }
        
        .btn-primary:active {
            transform: scale(0.98);
        }
        
        .footer {
            text-align: center;
            padding: 40px;
            color: #888888;
            font-size: 14px;
        }
    </style>
</head>
<body>
    <div class="header">
        <svg class="logo" viewBox="0 0 80 32" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="80" height="32" rx="4" fill="#00DC00"/>
            <text x="10" y="22" fill="white" font-family="Inter, sans-serif" font-weight="700" font-size="14">SAGE</text>
        </svg>
        <span class="header-text">SOW Generator</span>
    </div>
    
    <div class="container">
        <h1>Generate Statement of Work</h1>
        <p class="subtitle">Enter client details to generate a customized SOW document.</p>
        
        <div class="form-card">
            <form method="POST" action="/generate">
                <div class="form-group">
                    <label for="client_name">Client Name</label>
                    <input 
                        type="text" 
                        id="client_name"
                        name="client_name" 
                        required 
                        placeholder="Enter client name"
                    >
                </div>
                
                <div class="form-group">
                    <label for="num_leases">Number of Leases</label>
                    <input 
                        type="number" 
                        id="num_leases"
                        name="num_leases" 
                        required 
                        min="0" 
                        placeholder="Enter number of leases"
                    >
                </div>
                
                <button type="submit" class="btn-primary">Generate SOW</button>
            </form>
        </div>
    </div>
    
    <div class="footer">
        © 2024 Sage. All rights reserved.
    </div>
</body>
</html>
"""

RESULT_PAGE = """
<!DOCTYPE html>
<html>
<head>
    <title>SOW Ready | Sage</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            background: #F7F7F7;
            color: #1E1E1E;
            min-height: 100vh;
        }
        
        .header {
            background: #FFFFFF;
            padding: 20px 40px;
            border-bottom: 1px solid #E5E5E5;
            display: flex;
            align-items: center;
            gap: 12px;
        }
        
        .logo {
            width: 80px;
            height: auto;
        }
        
        .header-text {
            font-size: 24px;
            font-weight: 700;
            color: #1D6F42;
        }
        
        .container {
            max-width: 700px;
            margin: 60px auto;
            padding: 0 20px;
        }
        
        h1 {
            font-size: 32px;
            font-weight: 700;
            color: #1E1E1E;
            margin-bottom: 8px;
        }
        
        .subtitle {
            color: #666666;
            font-size: 16px;
            margin-bottom: 40px;
        }
        
        .info-card {
            background: #FFFFFF;
            padding: 24px;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
            margin-bottom: 24px;
            display: flex;
            gap: 40px;
        }
        
        .info-item {
            display: flex;
            flex-direction: column;
            gap: 4px;
        }
        
        .info-label {
            font-size: 12px;
            font-weight: 600;
            color: #888888;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }
        
        .info-value {
            font-size: 18px;
            font-weight: 600;
            color: #1E1E1E;
        }
        
        .sku-card {
            background: #FFFFFF;
            padding: 32px;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
            margin-bottom: 24px;
        }
        
        .sku-header {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 24px;
        }
        
        .sku-icon {
            width: 40px;
            height: 40px;
            background: #E8F5E8;
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
        }
        
        .sku-title {
            font-size: 20px;
            font-weight: 700;
            color: #1D6F42;
        }
        
        .sku-item {
            background: #F9F9F9;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 12px;
            border-left: 4px solid #00DC00;
        }
        
        .sku-item:last-child {
            margin-bottom: 0;
        }
        
        .sku-name {
            font-size: 16px;
            font-weight: 700;
            color: #1E1E1E;
            margin-bottom: 6px;
        }
        
        .sku-details {
            font-size: 14px;
            color: #666666;
        }
        
        .sku-price {
            font-weight: 600;
            color: #1D6F42;
        }
        
        .sku-quantity {
            display: inline-block;
            background: #00DC00;
            color: white;
            padding: 2px 10px;
            border-radius: 12px;
            font-size: 12px;
            font-weight: 600;
            margin-left: 8px;
        }
        
        .download-card {
            background: linear-gradient(135deg, #1D6F42 0%, #2A8B54 100%);
            padding: 40px;
            border-radius: 12px;
            text-align: center;
            color: white;
        }
        
        .download-title {
            font-size: 20px;
            font-weight: 700;
            margin-bottom: 8px;
        }
        
        .download-subtitle {
            font-size: 14px;
            opacity: 0.9;
            margin-bottom: 24px;
        }
        
        .btn-download {
            display: inline-flex;
            align-items: center;
            gap: 10px;
            background: #FFFFFF;
            color: #1D6F42;
            padding: 16px 32px;
            font-size: 16px;
            font-weight: 600;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            transition: transform 0.1s, box-shadow 0.2s;
            font-family: inherit;
        }
        
        .btn-download:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
        }
        
        .back-link {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            margin-top: 32px;
            color: #1D6F42;
            text-decoration: none;
            font-weight: 500;
            font-size: 14px;
        }
        
        .back-link:hover {
            text-decoration: underline;
        }
        
        .footer {
            text-align: center;
            padding: 40px;
            color: #888888;
            font-size: 14px;
        }
    </style>
</head>
<body>
    <div class="header">
        <svg class="logo" viewBox="0 0 80 32" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="80" height="32" rx="4" fill="#00DC00"/>
            <text x="10" y="22" fill="white" font-family="Inter, sans-serif" font-weight="700" font-size="14">SAGE</text>
        </svg>
        <span class="header-text">SOW Generator</span>
    </div>
    
    <div class="container">
        <h1>SOW Generated Successfully</h1>
        <p class="subtitle">Your Statement of Work is ready for download.</p>
        
        <div class="info-card">
            <div class="info-item">
                <span class="info-label">Client</span>
                <span class="info-value">{{ client_name }}</span>
            </div>
        </div>
        
        <div class="sku-card">
            <div class="sku-header">
                <div class="sku-icon">📦</div>
                <span class="sku-title">Recommended SKUs</span>
            </div>
            
            {% for sku in skus %}
            <div class="sku-item">
                <div class="sku-name">
                    {{ sku.sku }}
                    {% if sku.quantity and sku.quantity > 1 %}
                    <span class="sku-quantity">× {{ sku.quantity }}</span>
                    {% endif %}
                </div>
                <div class="sku-details">
                    {% if sku.hours %}
                        {{ sku.hours }} hours @ ${{ sku.rate }}/hr = <span class="sku-price">${{ "{:,}".format(sku.total) }}</span>
                    {% else %}
                        {{ sku.description }}
                    {% endif %}
                </div>
            </div>
            {% endfor %}
        </div>
        
        <div class="download-card">
            <div class="download-title">Your document is ready!</div>
            <div class="download-subtitle">Click below to download the PDF</div>
            <form method="GET" action="/download" style="display: inline;">
                <input type="hidden" name="filename" value="{{ filename }}">
                <button type="submit" class="btn-download">
                    <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                        <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
                        <polyline points="7 10 12 15 17 10"/>
                        <line x1="12" y1="15" x2="12" y2="3"/>
                    </svg>
                    Download PDF
                </button>
            </form>
        </div>
        
        <a href="/" class="back-link">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                <line x1="19" y1="12" x2="5" y2="12"/>
                <polyline points="12 19 5 12 12 5"/>
            </svg>
            Generate another SOW
        </a>
    </div>
    
    <div class="footer">
        © 2024 Sage. All rights reserved.
    </div>
</body>
</html>
"""

@app.route("/")
def home():
    return render_template_string(HTML_FORM)

@app.route("/generate", methods=["POST"])
def generate():
    client_name = request.form.get("client_name", "Client")
    num_leases = request.form.get("num_leases", "0")

    # Compute SKUs
    skus = compute_skus(num_leases)
    
    # Generate document
    doc = Document("template.docx")
    
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        if "{{client_name}}" in para.text:
            para.text = para.text.replace("{{client_name}}", client_name)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{client_name}}" in cell.text:
                    cell.text = cell.text.replace("{{client_name}}", client_name)

    # Save files
    safe_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).strip()
    docx_path = f"/tmp/{safe_name}_SOW.docx"
    pdf_path = f"/tmp/{safe_name}_SOW.pdf"

    doc.save(docx_path)

    # Convert to PDF
    try:
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", "/tmp", docx_path
        ], check=True)
    except subprocess.CalledProcessError as e:
        print("Error converting document:", e)
        return "Error generating PDF", 500
    
    # Return result page with SKUs
    return render_template_string(
        RESULT_PAGE,
        client_name=client_name,
        skus=skus,
        filename=f"{safe_name}_SOW.pdf"
    )

@app.route("/download")
def download():
    filename = request.args.get("filename", "SOW.pdf")
    filepath = f"/tmp/{filename}"

    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True, download_name=filename)
    else:
        return "File not found", 404

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
